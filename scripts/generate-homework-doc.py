from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SHOT_DIR = next(ROOT.rglob("01-login.png")).parent
OUT_DIR = SHOT_DIR.parent


def setup_doc():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    return doc


def read_text(rel_path):
    return (ROOT / rel_path).read_text(encoding="utf-8")


def add_code(doc, title, rel_path, markers=None, max_chars=2600):
    doc.add_heading(title, level=2)
    text = read_text(rel_path)
    if markers:
        start, end = markers
        start_index = text.find(start)
        end_index = text.find(end, start_index + len(start)) if end else -1
        if start_index >= 0:
            text = text[start_index:end_index + len(end)] if end_index > start_index else text[start_index:]
    if len(text) > max_chars:
        text = text[:max_chars] + "\n...（节选）"
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8)


def main():
    doc = setup_doc()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("0703 作业：物业管理系统 RBAC 权限体系实现")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Spring Boot + Vue 3 + Element Plus + MyBatis/MyBatis-Plus + PageHelper").bold = True

    doc.add_heading("一、作业完成内容", level=1)
    for item in [
        "完成后端框架优化：全局异常统一处理已保留，新增事务管理配置、OpenAPI 接口文档、MyBatis-Plus、PageHelper 分页。",
        "落地 RBAC 权限体系：新增角色表、权限表、用户角色关系表、角色权限关系表。",
        "实现登录认证：账号密码从数据库 sys_user 表读取，登录成功返回 token、角色、菜单和按钮权限。",
        "实现接口授权：后端新增 @RequirePermission，写操作按按钮权限码校验。",
        "实现用户管理、角色管理、菜单权限管理、用户角色分配、角色权限分配。",
        "实现按钮级权限控制：前端根据后端返回的 permission codes 隐藏无权限按钮。",
        "完成接口验证：登录、RBAC 当前权限、角色列表、权限列表、住户分页、OpenAPI 文档均可访问。",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("二、核心接口验证结果", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "验证项"
    table.rows[0].cells[1].text = "结果"
    for key, value in [
        ("登录接口", "POST /api/auth/login，返回 code=200，用户 admin"),
        ("角色数量", "4 个角色：系统管理员、物业客服主管、财务专员、工程维修组长"),
        ("权限数量", "26 个权限，包含菜单权限和按钮权限"),
        ("当前用户权限", "8 个菜单权限、18 个按钮权限"),
        ("分页接口", "GET /api/residents/page?pageNum=1&pageSize=3，total=6，records=3"),
        ("接口文档", "http://127.0.0.1:8080/swagger-ui/index.html 可访问"),
    ]:
        row = table.add_row().cells
        row[0].text = key
        row[1].text = value

    doc.add_heading("三、核心源代码节选", level=1)
    add_code(
        doc,
        "1. Maven 依赖：MyBatis-Plus、PageHelper、OpenAPI",
        "src/backend/pom.xml",
        ("<dependency>\n\t\t\t<groupId>com.baomidou</groupId>", "<artifactId>springdoc-openapi-starter-webmvc-ui</artifactId>\n\t\t\t<version>2.8.9</version>\n\t\t</dependency>"),
    )
    add_code(doc, "2. RBAC 数据库表结构", "src/backend/src/main/resources/schema.sql", ("CREATE TABLE sys_role", "CREATE TABLE resident"))
    add_code(doc, "3. 接口权限注解", "src/backend/src/main/java/com/training/management/common/annotation/RequirePermission.java")
    add_code(doc, "4. 登录返回用户权限", "src/backend/src/main/java/com/training/management/service/AuthService.java", ("public LoginUserVO buildLoginUser", "public String sha256"))
    add_code(doc, "5. RBAC 授权事务逻辑", "src/backend/src/main/java/com/training/management/service/RbacService.java", ("@Transactional\n    public void assignUserRoles", "private Map<String, Object> buildPermissionPayload"))
    add_code(doc, "6. RBAC 控制器接口", "src/backend/src/main/java/com/training/management/controller/RbacController.java", max_chars=3200)
    add_code(doc, "7. 前端按钮权限判断", "src/frontend/src/utils/roles.js", ("export function canAccessModule", None), max_chars=1800)

    doc.add_heading("四、系统功能截图", level=1)
    screenshots = [
        ("登录认证页面", "01-login.png"),
        ("首页仪表盘", "02-dashboard.png"),
        ("住户管理按钮级权限", "03-residents-button-permission.png"),
        ("RBAC 用户管理", "04-rbac-users.png"),
        ("RBAC 角色管理", "05-rbac-roles.png"),
        ("RBAC 菜单与按钮权限", "06-rbac-permissions.png"),
        ("RBAC 授权分配", "07-rbac-assign.png"),
        ("Swagger/OpenAPI 接口文档", "08-swagger.png"),
    ]
    for heading, filename in screenshots:
        doc.add_heading(heading, level=2)
        image_path = SHOT_DIR / filename
        if image_path.exists():
            doc.add_picture(str(image_path), width=Inches(6.4))
        else:
            doc.add_paragraph(f"截图缺失：{filename}")

    doc.add_heading("五、提交说明", level=1)
    doc.add_paragraph("邮件主题格式：0703作业+学号+姓名。Word 文档已整理核心源代码和系统功能截图，可按老师要求发送至指定邮箱并抄送。")

    out_path = OUT_DIR / "0703作业-RBAC权限体系实现.docx"
    doc.save(out_path)
    print(out_path)


if __name__ == "__main__":
    main()
